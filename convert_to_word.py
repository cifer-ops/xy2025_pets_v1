import docx
import markdown
import re
from docx.shared import Pt, Inches

# 读取markdown格式的文本文件
with open('论文.txt', 'r', encoding='utf-8') as file:
    content = file.read()

# 创建Word文档对象
doc = docx.Document()

# 设置文档样式
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# 处理标题和内容
lines = content.split('\n')
i = 0
while i < len(lines):
    line = lines[i].strip()
    
    # 处理一级标题 (#)
    if line.startswith('# '):
        title = line[2:].strip()
        doc.add_heading(title, level=0)
    
    # 处理二级标题 (##)
    elif line.startswith('## '):
        title = line[3:].strip()
        doc.add_heading(title, level=1)
    
    # 处理三级标题 (###)
    elif line.startswith('### '):
        title = line[4:].strip()
        doc.add_heading(title, level=2)
    
    # 处理四级标题 (####)
    elif line.startswith('#### '):
        title = line[5:].strip()
        doc.add_heading(title, level=3)
    
    # 处理列表项
    elif line.startswith('- ') or line.startswith('* '):
        text = line[2:].strip()
        p = doc.add_paragraph()
        p.add_run('• ' + text)
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 处理有序列表
    elif re.match(r'^\d+\.\s', line):
        text = re.sub(r'^\d+\.\s', '', line).strip()
        p = doc.add_paragraph()
        p.add_run(line)
        p.paragraph_format.left_indent = Inches(0.25)
    
    # 处理表格
    elif line.startswith('|') and i + 1 < len(lines) and lines[i+1].startswith('|'):
        # 获取表格头
        header_cells = [cell.strip() for cell in line.split('|')[1:-1]]
        
        # 跳过分隔行
        i += 1
        
        # 创建表格
        table = doc.add_table(rows=1, cols=len(header_cells))
        table.style = 'Table Grid'
        
        # 填充表头
        header_row = table.rows[0].cells
        for j, cell_text in enumerate(header_cells):
            header_row[j].text = cell_text
        
        # 填充数据行
        i += 1
        while i < len(lines) and lines[i].startswith('|'):
            row_cells = [cell.strip() for cell in lines[i].split('|')[1:-1]]
            if len(row_cells) > 0:
                row = table.add_row().cells
                for j, cell_text in enumerate(row_cells):
                    if j < len(row):
                        row[j].text = cell_text
            i += 1
        
        # 由于我们已经处理了当前行，需要减1以抵消后面的i+=1
        i -= 1
    
    # 处理普通段落
    elif line:
        doc.add_paragraph(line)
    
    # 处理空行，添加段落间距
    else:
        doc.add_paragraph()
    
    i += 1

# 保存文档
doc.save('宠物领养平台系统论文.docx')

print('转换完成，已保存为"宠物领养平台系统论文.docx"') 